"""DOCX exporters for chat, brief, and library conversations.

All three exporters return raw bytes (the v3.x base64 wrapper was a
PyWebView bridge constraint; FastAPI's StreamingResponse / Response handles
binary directly). Arabic RTL handling preserved from v3.x — `make_rtl` +
`set_cs_font` carry the hard-won OOXML knowledge about Word's schema
ordering and complex-script font handling.
"""

from __future__ import annotations

import html
import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.exports import (
    _COMPLIANCE_CHROME,
    _compliance_check_strings,
    format_dual_date,
    resolve_compliance_lang,
)

CHAT_ACCENT = (0x00, 0xD4, 0xAA)    # teal
LIB_ACCENT = (0x60, 0xA5, 0xFA)     # blue
GREY_DARK = (0x37, 0x41, 0x51)
GREY_MED = (0x6B, 0x72, 0x80)
GREY_LIGHT = (0x9C, 0xA3, 0xAF)

AR_FONT = "Noto Naskh Arabic"
EN_FONT = "Calibri"

_ARABIC_RE = re.compile(r"[؀-ۿ]")


def is_arabic(s: str) -> bool:
    return bool(_ARABIC_RE.search(s))


def _make_rtl(paragraph) -> None:
    """Insert <w:bidi/> at pPr position 0.

    Word's OOXML schema requires bidi BEFORE jc; appending after jc puts it
    out of order and Word silently ignores the flag (paragraph renders LTR
    despite RIGHT alignment). Insert at position 0 to guarantee correctness.
    """
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.insert(0, OxmlElement("w:bidi"))


def _set_cs_font(run, font_name: str = AR_FONT) -> None:
    """Apply complex-script font handling so Arabic renders in the chosen font.

    Without rFonts:cs Word falls back to Times New Roman for Arabic glyphs
    regardless of run.font.name. Also mirrors size and bold to their CS
    equivalents (szCs, bCs) and sets the run-level rtl + lang:bidi flags.
    """
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    sz_elem = rPr.find(qn("w:sz"))
    if sz_elem is not None:
        sz_val = sz_elem.get(qn("w:val"))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), sz_val)
    if rPr.find(qn("w:b")) is not None and rPr.find(qn("w:bCs")) is None:
        rPr.append(OxmlElement("w:bCs"))
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:bidi"), "ar-SA")
    if rPr.find(qn("w:rtl")) is None:
        rPr.append(OxmlElement("w:rtl"))


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _sanitize_md(text: str) -> str:
    """Strip markdown markup tokens for DOCX plain-run rendering."""
    text = html.unescape(text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[\-\*]\s+", "- ", text, flags=re.MULTILINE)
    return text


def _split_chat_blocks(content: str):
    """Yield ('para', text) or ('table', [[row], ...]).

    Detects markdown tables (pipe row followed by a separator row of dashes)
    so they render as native Word tables instead of raw `|col|col|` dumps.
    """
    lines = content.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        if line.startswith("|") and i + 1 < n:
            sep = lines[i + 1].strip()
            if sep.startswith("|") and re.match(r"^\|[\s\-:|]+\|?\s*$", sep):
                header = [c.strip() for c in line.strip().strip("|").split("|")]
                rows = [header]
                j = i + 2
                while j < n and lines[j].strip().startswith("|"):
                    rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                    j += 1
                yield "table", rows
                i = j
                continue
        stripped = line.strip()
        if stripped:
            yield "para", stripped
        i += 1


def _set_margins(doc, cm: float = 2.0) -> None:
    for section in doc.sections:
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)


def _add_paragraph_text(
    doc, text: str, *, size: int = 11, bold: bool = False, color=None, font=EN_FONT
):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ar = is_arabic(text)
    if ar:
        _make_rtl(p)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    if ar:
        _set_cs_font(r, AR_FONT)
    else:
        r.font.name = font
    return p, r


def _write_table_cell(cell, text: str, *, ar: bool, bold: bool = False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if ar:
        _make_rtl(p)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    if ar:
        _set_cs_font(r, AR_FONT)
    else:
        r.font.name = EN_FONT


def _bytes(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_chat_docx(messages: list[dict], date_format: str = "dual") -> bytes:
    doc = Document()
    _set_margins(doc, 2)

    _add_paragraph_text(doc, "TadqeeqAI Chat Export", size=20, bold=True, color=CHAT_ACCENT)
    _add_paragraph_text(
        doc,
        f"Generated: {format_dual_date(lang='en', mode=date_format, with_time=True)}",
        size=10,
        color=GREY_MED,
    )
    doc.add_paragraph()

    for msg in messages:
        role = "You" if msg.get("role") == "user" else "TadqeeqAI"
        _add_paragraph_text(doc, role, size=11, bold=True, color=CHAT_ACCENT)
        content = _sanitize_md(msg.get("content", ""))
        for kind, payload in _split_chat_blocks(content):
            if kind == "table":
                table_is_ar = any(is_arabic(" ".join(r)) for r in payload)
                wtable = doc.add_table(rows=len(payload), cols=len(payload[0]))
                wtable.alignment = (
                    WD_TABLE_ALIGNMENT.RIGHT if table_is_ar else WD_TABLE_ALIGNMENT.LEFT
                )
                wtable.style = "Table Grid"
                if table_is_ar:
                    tblPr = wtable._tbl.find(qn("w:tblPr"))
                    if tblPr is None:
                        tblPr = OxmlElement("w:tblPr")
                        wtable._tbl.insert(0, tblPr)
                    if tblPr.find(qn("w:bidiVisual")) is None:
                        tblPr.append(OxmlElement("w:bidiVisual"))
                for ri, row in enumerate(payload):
                    for ci, cell_text in enumerate(row):
                        cell = wtable.rows[ri].cells[ci]
                        if ri == 0:
                            _shade_cell(cell, "374151")
                            _write_table_cell(cell, cell_text, ar=table_is_ar, bold=True,
                                              color=(0xFF, 0xFF, 0xFF))
                        else:
                            _write_table_cell(cell, cell_text, ar=table_is_ar)
                doc.add_paragraph()
            else:
                _add_paragraph_text(doc, payload, size=11)

        if msg.get("sources"):
            sources_text = "Sources: " + ", ".join(
                s.get("article", "") for s in msg["sources"][:3]
            )
            _add_paragraph_text(doc, sources_text, size=9, color=GREY_MED)
        doc.add_paragraph()

    _add_paragraph_text(doc, "Generated by TadqeeqAI", size=8, color=GREY_LIGHT)
    return _bytes(doc)


def export_library_docx(
    messages: list[dict],
    category_label: str = "",
    clause_title: str = "",
    date_format: str = "dual",
) -> bytes:
    doc = Document()
    _set_margins(doc, 2)

    _add_paragraph_text(
        doc, "TadqeeqAI · Clause Library Discussion", size=20, bold=True, color=LIB_ACCENT
    )
    header_label = clause_title or category_label or "Clause Library"
    topic_p = doc.add_paragraph()
    topic_p.add_run("Topic: ").bold = True
    topic_p.add_run(header_label)
    if category_label and clause_title:
        cat_p = doc.add_paragraph()
        cat_p.add_run("Category: ").bold = True
        cat_p.add_run(category_label)
    _add_paragraph_text(
        doc,
        f"Generated: {format_dual_date(lang='en', mode=date_format, with_time=True)}",
        size=9,
        color=GREY_MED,
    )
    doc.add_paragraph()

    for msg in messages:
        label = "Question" if msg.get("role") == "user" else "Response"
        _add_paragraph_text(doc, label, size=12, bold=True, color=LIB_ACCENT)
        for line in msg.get("content", "").split("\n"):
            line = line.strip()
            if line:
                _add_paragraph_text(doc, line, size=11)
        if msg.get("sources"):
            sh_p = doc.add_paragraph()
            sh = sh_p.add_run("Sources:")
            sh.font.bold = True
            sh.font.size = Pt(10)
            for src in msg["sources"]:
                s_p = doc.add_paragraph()
                s_p.paragraph_format.left_indent = Cm(0.5)
                line = "• " + src.get("article", "")
                if src.get("title"):
                    line += f": {src['title']}"
                sr = s_p.add_run(line)
                sr.font.size = Pt(9)
                sr.font.color.rgb = RGBColor(*GREY_MED)
        doc.add_paragraph()

    _add_paragraph_text(
        doc, "Generated by TadqeeqAI · Clause Library", size=8, color=GREY_LIGHT
    )
    return _bytes(doc)


def export_compliance_docx(
    result: dict,
    date_format: str = "dual",
    lang: str = "auto",
) -> bytes:
    """Render a compliance result as DOCX with v3.2-style heading + per-check rows."""
    doc = Document()
    _set_margins(doc, 2.2)

    chrome_lang = resolve_compliance_lang(result, lang)
    chrome = _COMPLIANCE_CHROME[chrome_lang]
    score = int(result.get("score", 0))
    if score >= 80:
        score_color = (0x3F, 0xB9, 0x50)
    elif score >= 50:
        score_color = (0xFF, 0xBD, 0x2E)
    else:
        score_color = (0xFF, 0x45, 0x3A)

    _add_paragraph_text(doc, chrome["title"], size=20, bold=True, color=CHAT_ACCENT)
    _add_paragraph_text(doc, f"{chrome['doc_label']}: {result.get('filename', '')}", size=11)
    _add_paragraph_text(
        doc,
        f"{chrome['gen_label']}: "
        f"{format_dual_date(lang=chrome_lang, mode=date_format, with_time=True)}",
        size=10,
        color=GREY_MED,
    )
    _add_paragraph_text(
        doc,
        f"{chrome['score_word']}: {score}% ({chrome['score_label']})",
        size=14,
        bold=True,
        color=score_color,
    )
    doc.add_paragraph()

    for check in result.get("checks", []):
        is_pass = check.get("status") == "compliant"
        status_label = chrome["pass" if is_pass else "warn"]
        status_color = (0x3F, 0xB9, 0x50) if is_pass else (0xFF, 0xBD, 0x2E)
        icon = "✓" if is_pass else "⚠"
        strings = _compliance_check_strings(check, chrome_lang)
        _add_paragraph_text(
            doc,
            f"{icon} {strings['name']} ({status_label})",
            size=13,
            bold=True,
            color=status_color,
        )
        _add_paragraph_text(
            doc, f"{chrome['reg_label']}: {strings['regulation']}", size=10, color=GREY_MED
        )
        if strings["detail"]:
            _add_paragraph_text(doc, strings["detail"], size=11)
        doc.add_paragraph()

    _add_paragraph_text(doc, chrome["footer"], size=8, color=GREY_LIGHT)
    return _bytes(doc)


def export_brief_docx(text: str, date_format: str = "dual") -> bytes:
    """Render the executive brief as a multi-page DOCX, one page per language."""
    doc = Document()
    _set_margins(doc, 2.2)

    sections = [s.strip() for s in re.split(r"\n\s*---\s*\n", text) if s.strip()]
    if not sections:
        sections = [text or ""]

    for idx, section_text in enumerate(sections):
        section_is_ar = is_arabic(section_text)
        if idx > 0:
            pb = doc.add_paragraph()
            pb.add_run().add_break(WD_BREAK.PAGE)

        title = "الملخص التنفيذي" if section_is_ar else "Executive Brief"
        _add_paragraph_text(doc, title, size=20, bold=True, color=CHAT_ACCENT)
        _add_paragraph_text(
            doc,
            f"Generated: {format_dual_date(lang='en', mode=date_format, with_time=True)}",
            size=10,
            color=GREY_MED,
        )
        doc.add_paragraph()

        for raw in section_text.split("\n"):
            line = raw.rstrip()
            if not line.strip():
                continue
            if line.startswith("# "):
                _add_paragraph_text(doc, line[2:].strip(), size=18, bold=True, color=GREY_DARK)
            elif line.startswith("## "):
                _add_paragraph_text(doc, line[3:].strip(), size=14, bold=True, color=GREY_DARK)
            elif re.match(r"^[\-\*]\s+", line):
                body = re.sub(r"^[\-\*]\s+", "", line).strip()
                p, _ = _add_paragraph_text(doc, "• " + body, size=11)
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                _add_paragraph_text(doc, _sanitize_md(line), size=11)

    _add_paragraph_text(
        doc, "Generated by TadqeeqAI", size=8, color=GREY_LIGHT
    )
    return _bytes(doc)
