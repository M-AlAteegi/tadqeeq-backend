"""Document parsing + 6-point compliance checker.

Both classes are stateless utilities. Persistence happens in document_store.py.
Brief generation lives on TadqeeqRAG (it needs the embedder + provider).
"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO

import fitz
from docx import Document as DocxDocument

from app.config import settings


def detect_doc_language(text: str) -> str:
    arabic_chars = len(re.findall(r"[؀-ۿ]", text))
    return "ar" if arabic_chars > len(text) * 0.3 else "en"


def _quick_summary(text: str, filename: str) -> dict:
    """Cheap structural summary for the upload UI — no LLM needed."""
    text_lower = text.lower()
    has_arabic = bool(re.search(r"[؀-ۿ]", text))
    doc_type = "Document"
    if has_arabic:
        if any(k in text for k in ("صندوق", "صناديق الاستثمار", "مدير الصندوق", "وحدات الصندوق")):
            doc_type = "وثيقة صندوق استثماري"
        elif any(k in text for k in ("صكوك", "أدوات الدين", "سندات")):
            doc_type = "وثيقة صكوك / أدوات دين"
        elif any(k in text for k in ("ترخيص", "رخصة")):
            doc_type = "وثيقة ترخيص"
        elif any(k in text for k in ("عقد", "اتفاقية", "الأطراف")):
            doc_type = "عقد / اتفاقية"
        elif any(k in text for k in ("نشرة إصدار", "أوراق مالية")):
            doc_type = "نشرة إصدار"
        else:
            doc_type = "وثيقة تنظيمية"
    else:
        if any(k in text_lower for k in ("fund", "investment", "investor", "subscription")):
            doc_type = "Investment Fund Document"
        elif any(k in text_lower for k in ("sukuk", "bond", "debt instrument")):
            doc_type = "Sukuk/Debt Instrument Document"
        elif any(k in text_lower for k in ("license", "licensing", "authorization")):
            doc_type = "Licensing Document"
        elif any(k in text_lower for k in ("contract", "agreement", "party", "parties")):
            doc_type = "Contract/Agreement"
        elif any(k in text_lower for k in ("prospectus", "offering", "securities")):
            doc_type = "Securities Prospectus"
    return {"type": doc_type, "has_arabic": has_arabic, "word_count": len(text.split())}


class DocumentParseError(Exception):
    pass


class DocumentProcessor:
    """Stateless PDF/DOCX parser. Returns parsed text + structural metadata."""

    def parse(self, data: bytes, filename: str) -> dict:
        if len(data) > settings.max_upload_bytes:
            mb = len(data) / (1024 * 1024)
            raise DocumentParseError(
                f"File exceeds the {settings.max_upload_bytes // (1024*1024)} MB upload limit "
                f"({mb:.1f} MB received)."
            )
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext == "pdf":
            return self._parse_pdf(data, filename)
        if ext in ("docx", "doc"):
            return self._parse_docx(data, filename)
        raise DocumentParseError(f"Unsupported file type: .{ext}")

    def _parse_pdf(self, data: bytes, filename: str) -> dict:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
        except Exception as e:
            raise DocumentParseError(f"Failed to open PDF: {e}") from e
        try:
            page_count = doc.page_count
            if page_count > settings.max_document_pages:
                raise DocumentParseError(
                    f"Document too large. Maximum {settings.max_document_pages} pages allowed."
                )
            parts: list[str] = []
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    parts.append(f"[Page {page_num + 1}]\n{text}")
            full_text = "\n\n".join(parts)
        finally:
            doc.close()
        return {
            "filename": filename,
            "text": full_text,
            "page_count": page_count,
            "char_count": len(full_text),
            "summary": _quick_summary(full_text, filename),
        }

    def _parse_docx(self, data: bytes, filename: str) -> dict:
        try:
            doc = DocxDocument(BytesIO(data))
        except Exception as e:
            raise DocumentParseError(f"Failed to open DOCX: {e}") from e
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        full_text = "\n\n".join(parts)
        approx_pages = max(1, len(parts) // 20 + 1)
        return {
            "filename": filename,
            "text": full_text,
            "page_count": approx_pages,
            "char_count": len(full_text),
            "summary": _quick_summary(full_text, filename),
        }


_AR_DIACRITIC_RE = re.compile(r"[ً-ْٰـ]")
_AR_ALIF_RE = re.compile(r"[أإآٱ]")


def _normalize_arabic(text: str) -> str:
    """Tolerant normalization for Arabic keyword matching — strips tashkeel,
    collapses alif variants, normalizes alif-maksura / taa-marbuta."""
    if not text:
        return text
    text = _AR_DIACRITIC_RE.sub("", text)
    text = _AR_ALIF_RE.sub("ا", text)
    text = text.replace("ى", "ي").replace("ة", "ه")
    return text


def _matches_arabic_keyword(keyword: str, normalized_text: str) -> bool:
    """Match a normalized Arabic keyword allowing the optional ال prefix per word."""
    parts = [r"(?:ال)?" + re.escape(w) for w in keyword.split()]
    return bool(re.search(r"\s+".join(parts), normalized_text))


COMPLIANCE_CATEGORIES: dict[str, dict] = {
    "qualified_investor": {
        "name": "Qualified Investor Definition",
        "name_ar": "تعريف المستثمر المؤهل",
        "keywords": [
            "qualified investor", "accredited investor", "eligibility criteria",
            "مستثمر مؤهل", "المستثمر المؤهل", "مستثمرين مؤهلين", "مستثمر معتمد",
            "معايير الأهلية", "شروط الاهلية", "صافي الثروة", "الجدارة المالية",
        ],
        "regulation": "CMA Rules on Offer of Securities, Article 15",
        "regulation_ar": "قواعد طرح الأوراق المالية - هيئة السوق المالية، المادة 15",
        "description": "Documents offering securities must define qualified investor criteria",
        "description_ar": "يجب على مستندات طرح الأوراق المالية تعريف معايير المستثمر المؤهل",
    },
    "risk_disclosure": {
        "name": "Risk Disclosure",
        "name_ar": "الإفصاح عن المخاطر",
        "keywords": [
            "risk", "risks", "risk factors", "risk management", "market risk",
            "مخاطر", "المخاطر", "عوامل المخاطر", "مخاطر السوق", "إدارة المخاطر",
            "مخاطر الاستثمار", "مخاطر الائتمان", "مخاطر السيولة",
        ],
        "regulation": "CMA Rules on Offer of Securities, Article 22",
        "regulation_ar": "قواعد طرح الأوراق المالية - هيئة السوق المالية، المادة 22",
        "description": "Offering documents must include comprehensive risk disclosures",
        "description_ar": "يجب أن تتضمن مستندات الطرح إفصاحات شاملة عن المخاطر",
    },
    "capital_requirements": {
        "name": "Capital Requirements",
        "name_ar": "متطلبات رأس المال",
        "keywords": [
            "minimum capital", "paid-up capital", "capital adequacy", "share capital",
            "رأس المال", "رأسمال", "الحد الأدنى لرأس المال", "كفاية رأس المال",
            "رأس المال المدفوع", "رأس المال المصرح به",
        ],
        "regulation": "Finance Companies Control Law, Article 5",
        "regulation_ar": "نظام مراقبة شركات التمويل، المادة 5",
        "description": "Finance companies must meet minimum capital requirements",
        "description_ar": "يجب على شركات التمويل استيفاء الحد الأدنى لمتطلبات رأس المال",
    },
    "license_reference": {
        "name": "Licensing Information",
        "name_ar": "معلومات الترخيص",
        "keywords": [
            "license", "licensed", "licensing", "authorization", "CMA", "SAMA",
            "ترخيص", "مرخص", "مرخصة", "رخصة", "هيئة السوق المالية", "البنك المركزي السعودي",
            "مؤسسة النقد", "رقم الترخيص",
        ],
        "regulation": "Capital Market Law, Article 3",
        "regulation_ar": "نظام السوق المالية، المادة 3",
        "description": "Financial activities require proper licensing",
        "description_ar": "تتطلب الأنشطة المالية الحصول على ترخيص مناسب",
    },
    "fund_terms": {
        "name": "Fund Terms & Conditions",
        "name_ar": "شروط وأحكام الصناديق",
        "keywords": [
            "terms and conditions", "subscription", "management fee", "performance fee",
            "شروط وأحكام", "الشروط والأحكام", "رسوم الإدارة", "رسوم الاشتراك",
            "رسوم الأداء", "اتفاقية الاشتراك",
        ],
        "regulation": "Investment Funds Regulations, Article 20",
        "regulation_ar": "لائحة صناديق الاستثمار، المادة 20",
        "description": "Investment funds must clearly state terms and fees",
        "description_ar": "يجب على صناديق الاستثمار الإفصاح بوضوح عن الشروط والرسوم",
    },
    "disclosure_requirements": {
        "name": "Disclosure Requirements",
        "name_ar": "متطلبات الإفصاح",
        "keywords": [
            "disclosure", "material information", "transparency", "reporting obligation",
            "إفصاح", "الإفصاح", "متطلبات الإفصاح", "الإفصاح الكامل", "الإفصاح الدوري",
            "المعلومات الجوهرية", "الشفافية", "التقارير الدورية",
        ],
        "regulation": "CMA Rules on Offer of Securities, Article 30",
        "regulation_ar": "قواعد طرح الأوراق المالية - هيئة السوق المالية، المادة 30",
        "description": "Issuers must disclose all material information",
        "description_ar": "يجب على المُصدرين الإفصاح عن جميع المعلومات الجوهرية",
    },
}

CRITICAL_CATEGORY_IDS = {"qualified_investor", "risk_disclosure", "license_reference"}


def _localize(category: dict, status: str, found_keywords: list[str],
              pass_reason: str, language: str) -> dict:
    is_ar = language == "ar"
    name = category["name_ar"] if is_ar else category["name"]
    regulation = category["regulation_ar"] if is_ar else category["regulation"]
    description = category["description_ar"] if is_ar else category["description"]
    target_keywords = {
        kw.lower() for kw in category.get("keywords", [])
        if (re.search(r"[؀-ۿ]", kw) if is_ar else not re.search(r"[؀-ۿ]", kw))
    }
    same_lang_matches = [kw for kw in found_keywords if kw.lower() in target_keywords]
    if pass_reason == "matched":
        if same_lang_matches:
            listed = ", ".join(same_lang_matches[:3])
            detail = f"تم العثور على إشارات: {listed}" if is_ar else f"Found references: {listed}"
        else:
            detail = (
                f"تم اكتشاف إشارات متوافقة لـ{name}" if is_ar
                else f"Compliant references detected for {category['name'].lower()}"
            )
    elif pass_reason == "non_critical":
        detail = (
            "فئة غير حرجة — لم يتم الإبلاغ عنها وفق مستوى الصرامة الحالي." if is_ar
            else "Non-critical category — not flagged under current strictness."
        )
    else:
        detail = (
            f"يُنصح بإضافة معلومات حول {name}" if is_ar
            else f"Consider adding {category['name'].lower()} information"
        )
    return {"name": name, "regulation": regulation, "description": description, "detail": detail}


class ComplianceChecker:
    """6-point keyword-based compliance audit. Pure Python, no LLM."""

    def check(self, text: str, filename: str | None = None, strictness: str = "standard") -> dict:
        doc_language = detect_doc_language(text)
        text_lower = text.lower()
        normalized_text = _normalize_arabic(text)
        results: dict = {
            "filename": filename or "Document",
            "timestamp": datetime.now().isoformat(),
            "doc_language": doc_language,
            "checks": [],
            "summary": {"compliant": 0, "warnings": 0, "missing": 0},
        }
        for category_id, category in COMPLIANCE_CATEGORIES.items():
            found: list[str] = []
            for kw in category["keywords"]:
                if re.search(r"[؀-ۿ]", kw):
                    if _matches_arabic_keyword(_normalize_arabic(kw), normalized_text):
                        found.append(kw)
                elif kw.lower() in text_lower:
                    found.append(kw)
            if found:
                status, pass_reason = "compliant", "matched"
                results["summary"]["compliant"] += 1
            elif strictness == "critical_only" and category_id not in CRITICAL_CATEGORY_IDS:
                status, pass_reason = "compliant", "non_critical"
                results["summary"]["compliant"] += 1
            else:
                status, pass_reason = "warning", "missing"
                results["summary"]["warnings"] += 1
            check_record: dict = {
                "id": category_id,
                "status": status,
                "found_keywords": found,
                "pass_reason": pass_reason,
                "localized": {
                    "en": _localize(category, status, found, pass_reason, "en"),
                    "ar": _localize(category, status, found, pass_reason, "ar"),
                },
            }
            check_record.update(check_record["localized"][doc_language])
            results["checks"].append(check_record)
        total = results["summary"]["compliant"] + results["summary"]["warnings"]
        results["score"] = round((results["summary"]["compliant"] / total) * 100) if total > 0 else 100
        return results
